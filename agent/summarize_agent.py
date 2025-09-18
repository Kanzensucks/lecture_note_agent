from __future__ import annotations

import os
import logging
from typing import Optional, Iterable
from dataclasses import dataclass
from dotenv import load_dotenv
from docx import Document
import ollama

load_dotenv()

# ---------- Logging ----------
logger = logging.getLogger("lecture_note_agent")
if not logger.handlers:
    handler = logging.StreamHandler()
    fmt = logging.Formatter("[%(levelname)s] %(asctime)s %(name)s: %(message)s", "%Y-%m-%d %H:%M:%S")
    handler.setFormatter(fmt)
    logger.addHandler(handler)
logger.setLevel(logging.INFO)


# ---------- Config ----------
DEFAULT_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1")
MAX_CHARS_PER_CHUNK = int(os.getenv("MAX_CHARS_PER_CHUNK", "6000"))  # safe chunking


def load_prompt(prompt_path: str) -> str:
    """Load a prompt template from a text file."""
    with open(prompt_path, "r", encoding="utf-8") as f:
        return f.read().strip()


def _chunk_text(text: str, max_chars: int) -> Iterable[str]:
    text = text.strip()
    if len(text) <= max_chars:
        yield text
        return
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        # try to split on paragraph boundary
        split_at = text.rfind("\n\n", start, end)
        if split_at == -1 or split_at <= start + int(max_chars*0.4):
            split_at = end
        chunk = text[start:split_at].strip()
        if chunk:
            yield chunk
        start = split_at


@dataclass
class LLMResponse:
    content: str
    raw: dict


def _ollama_chat(prompt: str, model: str) -> LLMResponse:
    """Call Ollama chat API and return content."""
    logger.debug("Calling Ollama model=%s prompt_len=%d", model, len(prompt))
    res = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
    content = res.get("message", {}).get("content", "").strip()
    return LLMResponse(content=content, raw=res)


def summarize(text: str, prompt_template: str, *, model: Optional[str] = None) -> str:
    """Summarize a long text using a prompt template via Ollama.
    If the transcript is long, chunk it and then fuse the summaries.
    """
    model = model or DEFAULT_MODEL
    pieces = list(_chunk_text(text, MAX_CHARS_PER_CHUNK))
    logger.info("Summarizing %d chunk(s) with model=%s", len(pieces), model)

    if len(pieces) == 1:
        final_prompt = f"{prompt_template}\n\nTRANSCRIPT:\n{pieces[0]}"
        return _ollama_chat(final_prompt, model).content

    # Map phase: summarize each chunk
    interim_summaries = []
    map_template = (
        "You are helping summarize a long transcript in parts.\n"
        "Return a concise, self-contained summary of ONLY the part you are given.\n"
        "Focus on key facts, concepts, definitions, equations, and action items.\n"
        "Use clear headings and bullet points where natural.\n\n"
        "PART:\n{chunk}"
    )
    for i, chunk in enumerate(pieces, 1):
        logger.info("Summarizing chunk %d/%d (len=%d)", i, len(pieces), len(chunk))
        interim = _ollama_chat(map_template.format(chunk=chunk), model).content
        interim_summaries.append(interim)

    # Reduce phase: combine
    reduce_prompt = (
        f"{prompt_template}\n\n"
        "You are given several partial summaries from a long transcript.\n"
        "Fuse them into a single coherent output, removing duplication and keeping structure tight.\n"
        "Use clean markdown with short sections, bullet points, and numbered lists if appropriate.\n\n"
        "PARTIAL SUMMARIES:\n" + "\n\n----\n\n".join(interim_summaries)
    )
    combined = _ollama_chat(reduce_prompt, model).content
    return combined


def save_to_word(content: str, output_path: str) -> None:
    """Render simple markdown-ish content to a .docx file."""
    doc = Document()
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.strip().startswith(("- ", "* ")):
            doc.add_paragraph(line.strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
    logger.info("Saved Word file to %s", output_path)
